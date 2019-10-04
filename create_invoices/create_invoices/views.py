from datetime import datetime, date

from django.conf.global_settings import MEDIA_ROOT
from django.http import HttpResponse
from docx import *
from docx.shared import Inches


from django.contrib.auth.decorators import login_required
from django.shortcuts import redirect, render
from django.views import View
from django.contrib.auth import authenticate, login
from django.contrib.auth import logout as django_logout

from .models import Invoice
from .forms import InvoiceForm


class Login(View):
    def post(self, request):
        if request.user.is_anonymous:
            name = request.POST.get('name')
            password = request.POST.get('password')
            user = authenticate(username=name,
                                password=password)
            if user:
                login(request, user)
                return redirect('/invoices')
            else:
                template = "login.html"
                title = "Авторизация"
                return render(request, template, {'title': title, 'error': True})

        else:
            return redirect('/invoices')

    def get(self, request):
        if request.user.is_anonymous:
            template = "login.html"
            title = "Авторизация"
            return render(request, template, {'title': title})
        else:
            return redirect('/invoices')


@login_required
def logout(request):
    django_logout(request)
    return redirect('/login')


@login_required
def invoices(request):
    template = "invoices.html"
    invoices = Invoice.objects.all()
    title = "Список накладных"

    return render(request, template, {'invoices': invoices,
                                      'title': title})


class InvoiceDetail(View):
    def post(self, request, invoice_id):
        if request.user.is_anonymous:
            redirect('/login')
        invoice = Invoice.objects.filter(pk=invoice_id).first()
        if invoice:
            invoice.full_text = request.POST.get('full_text')
            invoice.text = request.POST.get('text')
            image_clear = request.POST.get('image-clear')
            if image_clear:
                invoice.image = None
            if request.FILES.get('image'):
               invoice.image = request.FILES.get('image')
            if request.POST.get('departure_date'):
                invoice.departure_date = datetime.strptime(request.POST.get('departure_date'), '%d.%m.%Y %H:%M:%S')
            else:
                invoice.departure_date = None
            if request.POST.get('receive_date'):
                invoice.receive_date = datetime.strptime(request.POST.get('receive_date'), '%d.%m.%Y %H:%M:%S')
            else:
                invoice.receive_date = None
            invoice.save()
        return redirect(request.path)

    def get(self, request, invoice_id):
        if request.user.is_anonymous:
            return redirect('/login')
        template = 'invoice_detail.html'
        invoice = Invoice.objects.filter(pk=invoice_id).first()
        title = invoice.name if invoice else 'Данный продукт не найден'
        return render(request, template, {'invoice': invoice,
                                          'title': title,
                                          'invoice_form': InvoiceForm(instance=invoice)})

@login_required
def redirect_to_invoices(request):
    return redirect('/invoices')


@login_required
def get_invoice_word(request):
    invoice = Invoice.objects.filter(pk=request.GET.get('id')).first()
    full_text = request.GET.get('full') == 'true'
    short_text = request.GET.get('short') == 'true'

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=test.docx'

    document = Document()
    document.add_heading(invoice.name, 0)
    if invoice.image:
        document.add_picture(invoice.image.path, width=Inches(1.25))

    document.add_paragraph('Дата создания документа ' + datetime.now().strftime('%d-%m-%Y %H:%M:%S'))

    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название поля'
    hdr_cells[1].text = 'Значение поля'

    records = []

    if full_text:
        if invoice.full_text:
            records.append(['Полное описание', invoice.full_text])

    if short_text:
        if invoice.text:
            records.append(['Краткое описание', invoice.text])

    if invoice.departure_date:
        records.append(['Дата отправления', invoice.departure_date.strftime('%d-%m-%Y %H:%M:%S')])

    if invoice.receive_date:
        records.append(['Дата получения', invoice.receive_date.strftime('%d-%m-%Y %H:%M:%S')])

    for key, value in records:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = value

    document.add_page_break()
    document.save(response)

    return response